"""Optional DOCX rendering of an attestation, for humans who want a
document.

render_docx() lays out a title page (attestation id, rollup), one table
per control (check, verdict, detail, evidence sha prefix), and a document
footer stating the only honest claim a rendering can make: the canonical
artifact is the attestation JSON. Byte stability is promised for the JSON
(attest/engine/report.py, tests/unit/test_determinism.py) and explicitly
NOT for this file -- python-docx embeds archive metadata this module does
not control.

python-docx is imported lazily inside render_docx(), so importing this
module costs nothing and works on machines without the package installed.

CLI wiring is future work: main.py deliberately does not yet grow a
`render` verb, and no render_cli() exists here. When the verb is added it
belongs with the runtime verbs (pure code, zero model spend).
"""

from __future__ import annotations

from pathlib import Path

from .engine import hash_obj

_FOOTER = (
    "The canonical artifact is the attestation JSON; this rendering is a "
    "convenience and carries no byte-stability claim."
)


def render_docx(attestation_body: dict, out_path: Path | str) -> Path:
    """Write a .docx rendering of one attestation body; returns the path.

    The attestation id is re-derived from the body -- the same id rule as
    attest/engine/report.py, 'att-' plus the first 16 hex characters of
    the canonical hash -- so the rendering cannot be captioned with an id
    its content does not hash to.
    """
    from docx import Document  # lazy: optional dependency (python-docx)

    attestation_id = "att-" + hash_obj(attestation_body)[:16]
    rollup = attestation_body["rollup"]
    counts = rollup["counts"]
    snapshot = attestation_body["snapshot"]

    doc = Document()

    # Title page ----------------------------------------------------------
    doc.add_heading(f"Attestation {attestation_id}", level=0)
    doc.add_paragraph(f"Rollup verdict: {rollup['verdict']}")
    doc.add_paragraph(
        f"Controls: {rollup['total']} "
        f"(pass {counts['pass']}, fail {counts['fail']}, unknown {counts['unknown']})"
    )
    doc.add_paragraph(f"Snapshot: {snapshot['id']} collected at {snapshot['collected_at']}")
    doc.add_paragraph(
        f"Engine {attestation_body['engine_version']}, "
        f"attestation schema {attestation_body['attest_version']}"
    )
    doc.add_page_break()

    # One table per control -----------------------------------------------
    for control in attestation_body["controls"]:
        doc.add_heading(f"{control['control_id']} -- {control['title']}", level=1)
        doc.add_paragraph(
            f"Verdict: {control['verdict']} "
            f"(raw {control['raw_verdict']}, on_unknown: {control['on_unknown']}) "
            f"-- spec {control['spec_sha256'][:16]}"
        )
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for cell, text in zip(table.rows[0].cells, ("check", "verdict", "detail", "evidence sha256")):
            cell.text = text
        for check in control["checks"]:
            sha = check["evidence"].get("sha256")
            row = table.add_row().cells
            row[0].text = check["check_id"]
            row[1].text = check["verdict"]
            row[2].text = check["detail"]
            row[3].text = sha[:16] if sha else "(no evidence)"

    # Footer ---------------------------------------------------------------
    doc.sections[0].footer.paragraphs[0].text = _FOOTER

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
